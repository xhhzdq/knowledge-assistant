"""Unit tests for DOCX paragraph and table extraction."""

from hashlib import sha256
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as OpenXmlDocument

from knowledge_assistant.exceptions import (
    DocumentParsingError,
    NoExtractableTextError,
    UnsupportedDocumentTypeError,
)
from knowledge_assistant.processing.parsers import DocxDocumentParser, select_parser

FIXTURE_DIR = Path(__file__).parent / "fixtures" / "week04"


def save_docx_to_bytes(document: OpenXmlDocument) -> bytes:
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_docx_fixture_extracts_title_and_paragraphs_in_order() -> None:
    content = (FIXTURE_DIR / "sample_document.docx").read_bytes()

    result = DocxDocumentParser().parse(content, "sample_document.docx")

    lines = result.pages[0].text.splitlines()
    assert lines == [
        "Week Four DOCX Fixture",
        "Structured document sample",
        "DOCX fixture body: headings and paragraphs become ordered text.",
        "A second paragraph checks that extraction preserves document order.",
    ]
    assert result.pages[0].page_number is None
    assert result.pages[0].requires_ocr is False
    assert result.content_hash == sha256(content).hexdigest()
    assert result.parser_name == "python-docx"
    assert result.parser_version


def test_docx_preserves_interleaved_paragraph_and_table_order() -> None:
    document = OpenXmlDocument()
    document.add_paragraph("before table")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Language"
    table.cell(1, 1).text = "Python"
    document.add_paragraph("after table")

    result = DocxDocumentParser().parse(save_docx_to_bytes(document), "table.docx")

    assert result.pages[0].text.splitlines() == [
        "before table",
        "Name\tValue",
        "Language\tPython",
        "after table",
    ]


def test_empty_docx_is_rejected() -> None:
    content = save_docx_to_bytes(OpenXmlDocument())

    with pytest.raises(NoExtractableTextError, match="no extractable text"):
        DocxDocumentParser().parse(content, "empty.docx")


def test_corrupt_docx_is_converted_to_parsing_error() -> None:
    with pytest.raises(DocumentParsingError, match="Unable to parse DOCX"):
        DocxDocumentParser().parse(b"not a zip package", "corrupt.docx")


def test_docx_parser_rejects_extension_mismatch() -> None:
    with pytest.raises(UnsupportedDocumentTypeError, match="does not support"):
        DocxDocumentParser().parse(b"content", "guide.pdf")


def test_parser_selector_can_select_docx_parser() -> None:
    parser = select_parser("GUIDE.DOCX", [DocxDocumentParser()])

    assert isinstance(parser, DocxDocumentParser)
